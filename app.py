from flask import Flask, render_template, request, redirect, url_for, session, flash, jsonify, make_response
import sqlite3
import os
import csv
import shutil
import json
from datetime import datetime, date
import calendar
from werkzeug.security import generate_password_hash, check_password_hash
from io import StringIO

app = Flask(__name__)
app.secret_key = 'your-secret-key-change-this-in-production'

# Database configuration
DATABASE = 'work_tracker.db'

def get_db_connection():
    conn = sqlite3.connect(DATABASE)
    conn.row_factory = sqlite3.Row
    return conn

def init_database():
    """Initialize the database with tables and sample data"""
    conn = get_db_connection()
    
    # Create tables
    conn.execute('''
        CREATE TABLE IF NOT EXISTS employees (
            emp_id TEXT PRIMARY KEY,
            first_name TEXT NOT NULL,
            last_name TEXT,
            username TEXT UNIQUE NOT NULL,
            email TEXT UNIQUE NOT NULL,
            department TEXT,
            role TEXT,
            password TEXT,
            is_admin INTEGER DEFAULT 0
        )
    ''')
    
    conn.execute('''
        CREATE TABLE IF NOT EXISTS projects (
            proj_id TEXT PRIMARY KEY,
            name TEXT NOT NULL,
            description TEXT,
            client TEXT,
            start_date DATE,
            end_date DATE,
            status TEXT DEFAULT 'Active',
            team_size INTEGER
        )
    ''')
    
    conn.execute('''
        CREATE TABLE IF NOT EXISTS work_progress (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            emp_id TEXT,
            proj_id TEXT,
            month TEXT,
            year INTEGER,
            percentage REAL,
            submission_date DATE,
            FOREIGN KEY (emp_id) REFERENCES employees (emp_id),
            FOREIGN KEY (proj_id) REFERENCES projects (proj_id),
            UNIQUE(emp_id, proj_id, month, year)
        )
    ''')
    
    # Insert sample employees with new password format: firstname@321 (add lastname first letter only if duplicates)
    employees_data = [
        ('EMP001', 'Shaji', 'Narayanan', 'shaji', 'shaji.narayanan@idroesse.com', 'Engineering', 'CAD Technician', generate_password_hash('shaji@321'), 0),
        ('EMP002', 'Shajimon', '', 'shajimon', 'shajimon@idroesse.com', 'Engineering', 'CAD Technician', generate_password_hash('shajimon@321'), 0),
        ('EMP003', 'Ibrahim', 'P.K', 'ibrahim', 'ibrahim.pk@idroesse.com', 'Engineering', 'CAD Technician', generate_password_hash('ibrahim@321'), 0),
        ('EMP004', 'Mohamed', 'Kaleefa', 'mohamed', 'mohamed.kaleefa@idroesse.com', 'Engineering', 'CAD Technician', generate_password_hash('mohamed@321'), 0),
        ('EMP005', 'Subeesh', 'Puthiyedath', 'subeesh', 'Subeesh.Puthiyedath@idroesse.com', 'Engineering', 'CAD Technician', generate_password_hash('subeesh@321'), 0),
        ('EMP006', 'Najeeb', 'E.K', 'najeeb', 'najeeb.ek@idroesse.com', 'Engineering', 'CAD Technician', generate_password_hash('najeeb@321'), 0),
        ('EMP007', 'Shakil', 'Ahmed', 'shakil', 'shakil.ahmed@idroesse.com', 'Engineering', 'CAD Technician', generate_password_hash('shakil@321'), 0),
        ('EMP008', 'Roy', 'Mathew', 'roy', 'roy.m@idroesse.com', 'Engineering', 'CAD Technician', generate_password_hash('roy@321'), 0),
        ('EMP009', 'Suhaib', 'Rahman', 'suhaib', 'suhaib.r@idroesse.com', 'Engineering', 'CAD Technician', generate_password_hash('suhaib@321'), 0),
        ('ADMIN001', 'Mahesh', 'Butani', 'mahesh.butani', 'mahesh.butani@idroesse.com', 'Management', 'Contracts Manager', generate_password_hash('m.butani@321'), 1)
    ]
    
    for emp in employees_data:
        try:
            conn.execute('''
                INSERT OR IGNORE INTO employees 
                (emp_id, first_name, last_name, username, email, department, role, password, is_admin) 
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', emp)
        except:
            pass
    
    # Insert sample projects
    projects_data = [
        ('PROJ201', 'AWSC - Commercial Spaces', 'Commercial spaces development project', 'AWSC', '2024-01-01', '2025-06-30', 'Active', 10),
        ('PROJ294', 'ADM - Stormwater Drainage Networks in Unserved Areas', 'Stormwater drainage infrastructure in unserved areas', 'ADM', '2024-02-01', '2025-08-31', 'Active', 10),
        ('PROJ286', 'Aldar - Dgn of Bida Al Mutawaa to Um Al Ashtan Road & Streetlighting', 'Road design and streetlighting from Bida Al Mutawaa to Um Al Ashtan', 'Aldar', '2024-01-15', '2025-05-30', 'Active', 10),
        ('PROJ208', 'MOEI - Road btwn MBZ City & KBZ (E84)-Fujairah', 'Road connection between MBZ City and KBZ (E84) in Fujairah', 'MOEI', '2023-11-01', '2024-12-31', 'Active', 10),
        ('PROJ298', 'MOEI - Dgn of TSE in Ajman', 'Design of Treated Sewage Effluent system in Ajman', 'MOEI', '2024-03-01', '2025-09-30', 'Active', 10),
        ('PROJ197', 'MOEI - Cont\'n of E87 Road to E99 - Fuj', 'Continuation of E87 Road to E99 in Fujairah', 'MOEI', '2023-10-01', '2024-11-30', 'Active', 10),
        ('PROJ302', 'ADSC - Abu Dhabi Sport Council HQ', 'Abu Dhabi Sport Council Headquarters development', 'ADSC', '2024-04-01', '2025-12-31', 'Active', 10),
        ('PROJ297', 'ADSC - Community Sports Facilities in Madinat Zayed & Sila\'a', 'Community sports facilities development in Madinat Zayed and Sila\'a', 'ADSC', '2024-03-15', '2025-10-31', 'Active', 10),
        ('PROJ303', 'HEC - Falcon Eye', 'Falcon Eye project development', 'HEC', '2024-04-15', '2025-11-30', 'Active', 10),
        ('PROJ309', 'AACM - Design of Northern Ring Rd Ext. & Al Ain-Sidira', 'Design of Northern Ring Road Extension and Al Ain-Sidira connection', 'AACM', '2024-05-01', '2026-02-28', 'Active', 10),
        ('PROJ314', 'Makan - Update the Infrastructure Design of B.A.S. Development Detailed Master Plan', 'Infrastructure design update for B.A.S. Development master plan', 'Makan', '2024-06-01', '2025-12-31', 'Active', 10),
        ('PROJ319', 'Ajman - Design of Sh. Ammar, Hamidiya & Jeddah Streets', 'Street design for Sh. Ammar, Hamidiya and Jeddah Streets in Ajman', 'Ajman Municipality', '2024-07-01', '2025-06-30', 'Active', 10),
        ('PROJ320', 'ADPIC-DMT - Design Review & Development (T2)', 'Design review and development for ADPIC-DMT Terminal 2', 'ADPIC-DMT', '2024-07-15', '2025-08-31', 'Active', 10),
        ('PROJ000', 'SRTA - Various', 'Various projects under SRTA', 'SRTA', '2024-01-01', '2025-12-31', 'Active', 10)
    ]
    
    for proj in projects_data:
        try:
            conn.execute('''
                INSERT OR IGNORE INTO projects 
                (proj_id, name, description, client, start_date, end_date, status, team_size) 
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            ''', proj)
        except:
            pass
    
    conn.commit()
    conn.close()

# Routes
@app.route('/')
def launcher():
    return render_template('launcher.html')

@app.route('/admin/login', methods=['GET', 'POST'])
def admin_login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        
        conn = get_db_connection()
        user = conn.execute('SELECT * FROM employees WHERE username = ? AND is_admin = 1', (username,)).fetchone()
        conn.close()
        
        if user and check_password_hash(user['password'], password):
            session['user_id'] = user['emp_id']
            session['user_name'] = f"{user['first_name']} {user['last_name']}"
            session['is_admin'] = user['is_admin']
            return redirect(url_for('admin_dashboard'))
        else:
            flash('Invalid username or password')
    
    return render_template('admin_login.html')

@app.route('/employee/login', methods=['GET', 'POST'])
def employee_login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        
        conn = get_db_connection()
        user = conn.execute('SELECT * FROM employees WHERE username = ? AND is_admin = 0', (username,)).fetchone()
        conn.close()
        
        if user and check_password_hash(user['password'], password):
            session['user_id'] = user['emp_id']
            session['user_name'] = f"{user['first_name']} {user['last_name']}"
            session['is_admin'] = user['is_admin']
            return redirect(url_for('employee_dashboard'))
        else:
            flash('Invalid username or password')
    
    # Get all employee usernames for dropdown
    conn = get_db_connection()
    employees = conn.execute('SELECT username, first_name, last_name FROM employees WHERE is_admin = 0 ORDER BY first_name').fetchall()
    conn.close()
    
    return render_template('employee_login.html', employees=employees)

@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('launcher'))

@app.route('/admin/dashboard')
def admin_dashboard():
    if not session.get('is_admin'):
        return redirect(url_for('admin_login'))
    
    # Get selected month and year from request parameters or default to previous month
    selected_month = request.args.get('month')
    selected_year = request.args.get('year', type=int)
    
    current_date = datetime.now()
    
    # Default to previous month if not specified
    if not selected_month or not selected_year:
        if current_date.month == 1:
            selected_month = 'December'
            selected_year = current_date.year - 1
        else:
            previous_month_num = current_date.month - 1
            selected_month = calendar.month_name[previous_month_num]
            selected_year = current_date.year
    
    conn = get_db_connection()
    
    # 1. EMPLOYEE SUMMARY - Performance by employee for selected period
    employee_summary = conn.execute('''
        SELECT e.emp_id, e.first_name, e.last_name, e.department,
               COUNT(DISTINCT wp.proj_id) as projects_worked,
               COUNT(wp.id) as total_entries,
               COALESCE(SUM(wp.percentage), 0) as total_percentage,
               COALESCE(AVG(wp.percentage), 0) as avg_percentage_per_project,
               MAX(wp.submission_date) as last_submission
        FROM employees e
        LEFT JOIN work_progress wp ON e.emp_id = wp.emp_id 
                                   AND wp.month = ? AND wp.year = ?
        WHERE e.is_admin = 0
        GROUP BY e.emp_id, e.first_name, e.last_name, e.department
        ORDER BY total_percentage DESC, e.first_name
    ''', (selected_month, selected_year)).fetchall()
    
    # 2. PROJECT SUMMARY - Utilization by project for selected period
    project_summary = conn.execute('''
        SELECT p.proj_id, p.name, p.client, p.status,
               COUNT(DISTINCT wp.emp_id) as employees_worked,
               COUNT(wp.id) as total_entries,
               COALESCE(SUM(wp.percentage), 0) as total_work_percentage,
               COALESCE(AVG(wp.percentage), 0) as avg_percentage_per_employee,
               ROUND(COALESCE(SUM(wp.percentage), 0) / 100.0, 2) as equivalent_employees
        FROM projects p
        LEFT JOIN work_progress wp ON p.proj_id = wp.proj_id 
                                   AND wp.month = ? AND wp.year = ?
                                   AND wp.percentage > 0
        GROUP BY p.proj_id, p.name, p.client, p.status
        ORDER BY total_work_percentage DESC, p.name
    ''', (selected_month, selected_year)).fetchall()
    
    # 3. MONTHLY SUMMARY - Performance metrics for selected month
    monthly_summary = conn.execute('''
        SELECT 
            COUNT(DISTINCT wp.emp_id) as employees_submitted,
            COUNT(DISTINCT wp.proj_id) as projects_with_work,
            COUNT(wp.id) as total_submissions,
            COALESCE(AVG(wp.percentage), 0) as avg_percentage_per_entry,
            COALESCE(SUM(wp.percentage), 0) as total_percentage_submitted,
            MIN(wp.submission_date) as first_submission,
            MAX(wp.submission_date) as last_submission
        FROM work_progress wp
        WHERE wp.month = ? AND wp.year = ? AND wp.percentage > 0
    ''', (selected_month, selected_year)).fetchone()
    
    # 4. YEARLY SUMMARY - Cumulative performance for selected year
    yearly_summary = conn.execute('''
        SELECT 
            COUNT(DISTINCT wp.emp_id) as unique_employees,
            COUNT(DISTINCT wp.proj_id) as unique_projects,
            COUNT(DISTINCT wp.month) as months_with_data,
            COUNT(wp.id) as total_yearly_submissions,
            COALESCE(AVG(wp.percentage), 0) as avg_yearly_percentage,
            COALESCE(SUM(wp.percentage), 0) as total_yearly_percentage
        FROM work_progress wp
        WHERE wp.year = ? AND wp.percentage > 0
    ''', (selected_year,)).fetchone()
    
    # Chart data for employee completion status (handle None values safely)
    employees_completed_list = [emp for emp in employee_summary if emp['total_percentage'] is not None and emp['total_percentage'] == 100]
    employees_partial_list = [emp for emp in employee_summary if emp['total_percentage'] is not None and 0 < emp['total_percentage'] < 100]
    employees_not_started_list = [emp for emp in employee_summary if emp['total_percentage'] is None or emp['total_percentage'] == 0]
    
    # Counts for charts
    employees_completed = len(employees_completed_list)
    employees_partial = len(employees_partial_list)
    employees_not_started = len(employees_not_started_list)
    
    # Chart data for project utilization (top 10 projects)
    top_projects = [proj for proj in project_summary if proj['total_work_percentage'] > 0][:10]
    
    # Monthly trend data for the current year
    monthly_trends = conn.execute('''
        SELECT month, 
               COUNT(DISTINCT emp_id) as employees,
               COUNT(DISTINCT proj_id) as projects,
               COUNT(id) as submissions,
               SUM(percentage) as total_work_percentage
        FROM work_progress 
        WHERE year = ? AND percentage > 0
        GROUP BY month
        ORDER BY 
        CASE month 
            WHEN 'January' THEN 1 WHEN 'February' THEN 2 WHEN 'March' THEN 3
            WHEN 'April' THEN 4 WHEN 'May' THEN 5 WHEN 'June' THEN 6
            WHEN 'July' THEN 7 WHEN 'August' THEN 8 WHEN 'September' THEN 9
            WHEN 'October' THEN 10 WHEN 'November' THEN 11 WHEN 'December' THEN 12
        END
    ''', (selected_year,)).fetchall()
    
    # Get available months and years for dropdown
    available_periods = conn.execute('''
        SELECT DISTINCT month, year 
        FROM work_progress 
        ORDER BY year DESC, 
        CASE month 
            WHEN 'January' THEN 1 WHEN 'February' THEN 2 WHEN 'March' THEN 3
            WHEN 'April' THEN 4 WHEN 'May' THEN 5 WHEN 'June' THEN 6
            WHEN 'July' THEN 7 WHEN 'August' THEN 8 WHEN 'September' THEN 9
            WHEN 'October' THEN 10 WHEN 'November' THEN 11 WHEN 'December' THEN 12
        END DESC
    ''').fetchall()
    
    # Recent submissions for the selected period
    recent_submissions = conn.execute('''
        SELECT wp.*, e.first_name, e.last_name, p.name as project_name
        FROM work_progress wp
        JOIN employees e ON wp.emp_id = e.emp_id
        JOIN projects p ON wp.proj_id = p.proj_id
        WHERE wp.month = ? AND wp.year = ?
        ORDER BY wp.submission_date DESC, e.first_name
        LIMIT 15
    ''', (selected_month, selected_year)).fetchall()
    
    conn.close()
    
    return render_template('admin_dashboard.html', 
                         # Summaries
                         employee_summary=employee_summary,
                         project_summary=project_summary,
                         monthly_summary=monthly_summary,
                         yearly_summary=yearly_summary,
                         
                         # Chart data - counts
                         employees_completed=employees_completed,
                         employees_partial=employees_partial,
                         employees_not_started=employees_not_started,
                         
                         # Chart data - actual lists
                         employees_completed_list=employees_completed_list,
                         employees_partial_list=employees_partial_list,
                         employees_not_started_list=employees_not_started_list,
                         
                         top_projects=top_projects,
                         monthly_trends=monthly_trends,
                         
                         # Filter data
                         selected_month=selected_month,
                         selected_year=selected_year,
                         available_periods=available_periods,
                         recent_submissions=recent_submissions,
                         
                         # Legacy compatibility
                         total_projects=len([p for p in project_summary if p['status'] == 'Active']),
                         total_employees=len([e for e in employee_summary if not e['total_percentage'] is None]),
                         previous_month_submissions=recent_submissions)

@app.route('/employee/dashboard')
def employee_dashboard():
    if not session.get('user_id') or session.get('is_admin'):
        return redirect(url_for('employee_login'))
    
    # Get previous month and year
    current_date = datetime.now()
    if current_date.month == 1:
        previous_month = 'December'
        previous_year = current_date.year - 1
    else:
        previous_month_num = current_date.month - 1
        previous_month = calendar.month_name[previous_month_num]
        previous_year = current_date.year
    
    conn = get_db_connection()
    
    # Get all active projects
    projects = conn.execute('SELECT * FROM projects WHERE status = "Active"').fetchall()
    
    # Check if employee has already submitted data for the previous month
    existing_data = conn.execute('''
        SELECT COUNT(*) as count FROM work_progress 
        WHERE emp_id = ? AND month = ? AND year = ?
    ''', (session['user_id'], previous_month, previous_year)).fetchone()
    
    has_submitted_data = existing_data['count'] > 0
    
    # Get current user's progress for previous month
    progress = {}
    total_percentage = 0
    for project in projects:
        result = conn.execute('''
            SELECT percentage FROM work_progress 
            WHERE emp_id = ? AND proj_id = ? AND month = ? AND year = ?
        ''', (session['user_id'], project['proj_id'], previous_month, previous_year)).fetchone()
        
        if result:
            progress[project['proj_id']] = result['percentage']
            total_percentage += result['percentage']
        else:
            progress[project['proj_id']] = 0
    
    # Get submission date if data exists
    submission_date = None
    if has_submitted_data:
        submission_result = conn.execute('''
            SELECT submission_date FROM work_progress 
            WHERE emp_id = ? AND month = ? AND year = ?
            ORDER BY submission_date DESC LIMIT 1
        ''', (session['user_id'], previous_month, previous_year)).fetchone()
        if submission_result:
            submission_date = submission_result['submission_date']
    
    conn.close()
    
    return render_template('employee_dashboard.html', 
                         projects=projects, 
                         progress=progress,
                         current_month=previous_month,
                         current_year=previous_year,
                         has_submitted_data=has_submitted_data,
                         total_percentage=total_percentage,
                         submission_date=submission_date)

@app.route('/employee/submit_progress', methods=['POST'])
def submit_progress():
    if not session.get('user_id') or session.get('is_admin'):
        return redirect(url_for('employee_login'))
    
    # Get previous month and year
    current_date = datetime.now()
    if current_date.month == 1:
        previous_month = 'December'
        previous_year = current_date.year - 1
    else:
        previous_month_num = current_date.month - 1
        previous_month = calendar.month_name[previous_month_num]
        previous_year = current_date.year
    
    conn = get_db_connection()
    
    # Check if employee has already submitted data for the previous month
    existing_data = conn.execute('''
        SELECT COUNT(*) as count FROM work_progress 
        WHERE emp_id = ? AND month = ? AND year = ?
    ''', (session['user_id'], previous_month, previous_year)).fetchone()
    
    if existing_data['count'] > 0:
        conn.close()
        flash('You have already submitted your work progress for this month. No further submissions allowed.', 'error')
        return redirect(url_for('employee_dashboard'))
    
    # Calculate total percentage
    total_percentage = 0
    progress_data = {}
    for key, value in request.form.items():
        if key.startswith('progress_'):
            proj_id = key.replace('progress_', '')
            percentage = float(value) if value else 0
            progress_data[proj_id] = percentage
            total_percentage += percentage
    
    # Validate that total is exactly 100%
    if abs(total_percentage - 100.0) > 0.01:  # Allow small floating point differences
        conn.close()
        flash(f'Total work percentage must equal 100%. Current total: {total_percentage:.1f}%', 'error')
        return redirect(url_for('employee_dashboard'))
    
    # Save only non-zero work percentage data
    for proj_id, percentage in progress_data.items():
        if percentage > 0:  # Only store projects with actual work done
            conn.execute('''
                INSERT INTO work_progress 
                (emp_id, proj_id, month, year, percentage, submission_date) 
                VALUES (?, ?, ?, ?, ?, ?)
            ''', (session['user_id'], proj_id, previous_month, previous_year, percentage, date.today()))
    
    conn.commit()
    conn.close()
    
    flash('Work percentage submitted successfully! Total: 100%. You cannot submit again for this month.', 'success')
    return redirect(url_for('employee_dashboard'))

@app.route('/admin/projects')
def admin_projects():
    if not session.get('is_admin'):
        return redirect(url_for('admin_login'))
    
    conn = get_db_connection()
    projects = conn.execute('SELECT * FROM projects ORDER BY start_date DESC').fetchall()
    conn.close()
    
    return render_template('admin_projects.html', projects=projects)

@app.route('/admin/projects/add', methods=['POST'])
def add_project():
    if not session.get('is_admin'):
        return redirect(url_for('admin_login'))
    
    try:
        proj_id = request.form['proj_id']
        name = request.form['name']
        description = request.form['description']
        client = request.form['client']
        start_date = request.form['start_date']
        end_date = request.form['end_date']
        team_size = int(request.form['team_size'])
        
        conn = get_db_connection()
        conn.execute('''
            INSERT INTO projects (proj_id, name, description, client, start_date, end_date, status, team_size)
            VALUES (?, ?, ?, ?, ?, ?, 'Active', ?)
        ''', (proj_id, name, description, client, start_date, end_date, team_size))
        conn.commit()
        conn.close()
        
        flash('Project added successfully!', 'success')
    except Exception as e:
        flash(f'Error adding project: {str(e)}', 'error')
    
    return redirect(url_for('admin_projects'))

@app.route('/admin/projects/edit/<proj_id>', methods=['POST'])
def edit_project(proj_id):
    if not session.get('is_admin'):
        return redirect(url_for('admin_login'))
    
    try:
        name = request.form['name']
        description = request.form['description']
        client = request.form['client']
        start_date = request.form['start_date']
        end_date = request.form['end_date']
        status = request.form['status']
        team_size = int(request.form['team_size'])
        
        conn = get_db_connection()
        conn.execute('''
            UPDATE projects SET name=?, description=?, client=?, start_date=?, end_date=?, status=?, team_size=?
            WHERE proj_id=?
        ''', (name, description, client, start_date, end_date, status, team_size, proj_id))
        conn.commit()
        conn.close()
        
        flash('Project updated successfully!', 'success')
    except Exception as e:
        flash(f'Error updating project: {str(e)}', 'error')
    
    return redirect(url_for('admin_projects'))

@app.route('/admin/projects/delete/<proj_id>', methods=['POST'])
def delete_project(proj_id):
    if not session.get('is_admin'):
        return redirect(url_for('admin_login'))
    
    try:
        conn = get_db_connection()
        # Check if project has any progress data
        progress_count = conn.execute('SELECT COUNT(*) as count FROM work_progress WHERE proj_id = ?', (proj_id,)).fetchone()['count']
        
        if progress_count > 0:
            # Don't delete, just mark as inactive
            conn.execute('UPDATE projects SET status = "Inactive" WHERE proj_id = ?', (proj_id,))
            flash('Project marked as inactive (has progress data).', 'warning')
        else:
            # Safe to delete
            conn.execute('DELETE FROM projects WHERE proj_id = ?', (proj_id,))
            flash('Project deleted successfully!', 'success')
        
        conn.commit()
        conn.close()
    except Exception as e:
        flash(f'Error deleting project: {str(e)}', 'error')
    
    return redirect(url_for('admin_projects'))

@app.route('/admin/employees')
def admin_employees():
    if not session.get('is_admin'):
        return redirect(url_for('admin_login'))
    
    conn = get_db_connection()
    employees = conn.execute('SELECT * FROM employees ORDER BY first_name').fetchall()
    conn.close()
    
    return render_template('admin_employees.html', employees=employees)

@app.route('/admin/employees/add', methods=['POST'])
def add_employee():
    if not session.get('is_admin'):
        return redirect(url_for('admin_login'))
    
    try:
        emp_id = request.form['emp_id']
        first_name = request.form['first_name']
        last_name = request.form['last_name']
        username = request.form['username']
        email = request.form['email']
        department = request.form['department']
        role = request.form['role']
        password = request.form['password']
        is_admin = 1 if request.form.get('is_admin') else 0
        
        hashed_password = generate_password_hash(password)
        
        conn = get_db_connection()
        conn.execute('''
            INSERT INTO employees (emp_id, first_name, last_name, username, email, department, role, password, is_admin)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (emp_id, first_name, last_name, username, email, department, role, hashed_password, is_admin))
        conn.commit()
        conn.close()
        
        flash('Employee added successfully!', 'success')
    except Exception as e:
        flash(f'Error adding employee: {str(e)}', 'error')
    
    return redirect(url_for('admin_employees'))

@app.route('/admin/employees/edit/<emp_id>', methods=['POST'])
def edit_employee(emp_id):
    if not session.get('is_admin'):
        return redirect(url_for('admin_login'))
    
    try:
        first_name = request.form['first_name']
        last_name = request.form['last_name']
        username = request.form['username']
        email = request.form['email']
        department = request.form['department']
        role = request.form['role']
        is_admin = 1 if request.form.get('is_admin') else 0
        
        conn = get_db_connection()
        
        # Update password only if provided
        if request.form.get('password'):
            password = request.form['password']
            hashed_password = generate_password_hash(password)
            conn.execute('''
                UPDATE employees SET first_name=?, last_name=?, username=?, email=?, department=?, role=?, password=?, is_admin=?
                WHERE emp_id=?
            ''', (first_name, last_name, username, email, department, role, hashed_password, is_admin, emp_id))
        else:
            conn.execute('''
                UPDATE employees SET first_name=?, last_name=?, username=?, email=?, department=?, role=?, is_admin=?
                WHERE emp_id=?
            ''', (first_name, last_name, username, email, department, role, is_admin, emp_id))
        
        conn.commit()
        conn.close()
        
        flash('Employee updated successfully!', 'success')
    except Exception as e:
        flash(f'Error updating employee: {str(e)}', 'error')
    
    return redirect(url_for('admin_employees'))

@app.route('/admin/employees/delete/<emp_id>', methods=['POST'])
def delete_employee(emp_id):
    if not session.get('is_admin'):
        return redirect(url_for('admin_login'))
    
    try:
        conn = get_db_connection()
        # Check if employee has any progress data
        progress_count = conn.execute('SELECT COUNT(*) as count FROM work_progress WHERE emp_id = ?', (emp_id,)).fetchone()['count']
        
        if progress_count > 0:
            flash('Cannot delete employee with existing progress data.', 'error')
        else:
            # Safe to delete
            conn.execute('DELETE FROM employees WHERE emp_id = ?', (emp_id,))
            conn.commit()
            flash('Employee deleted successfully!', 'success')
        
        conn.close()
    except Exception as e:
        flash(f'Error deleting employee: {str(e)}', 'error')
    
    return redirect(url_for('admin_employees'))

@app.route('/admin/reports')
def admin_reports():
    if not session.get('is_admin'):
        return redirect(url_for('admin_login'))
    
    conn = get_db_connection()
    
    # Get progress data for reports with more detailed information
    progress_data = conn.execute('''
        SELECT wp.*, 
               e.first_name, e.last_name, e.email, e.department,
               p.name as project_name, p.client, p.description as project_description
        FROM work_progress wp
        JOIN employees e ON wp.emp_id = e.emp_id
        JOIN projects p ON wp.proj_id = p.proj_id
        ORDER BY wp.year DESC, 
        CASE wp.month 
            WHEN 'January' THEN 1 WHEN 'February' THEN 2 WHEN 'March' THEN 3
            WHEN 'April' THEN 4 WHEN 'May' THEN 5 WHEN 'June' THEN 6
            WHEN 'July' THEN 7 WHEN 'August' THEN 8 WHEN 'September' THEN 9
            WHEN 'October' THEN 10 WHEN 'November' THEN 11 WHEN 'December' THEN 12
        END DESC, 
        e.first_name, p.name
    ''').fetchall()
    
    # Get unique months/years for filtering
    months_years = conn.execute('''
        SELECT DISTINCT month, year FROM work_progress 
        ORDER BY year DESC, 
        CASE month 
            WHEN 'January' THEN 1 WHEN 'February' THEN 2 WHEN 'March' THEN 3
            WHEN 'April' THEN 4 WHEN 'May' THEN 5 WHEN 'June' THEN 6
            WHEN 'July' THEN 7 WHEN 'August' THEN 8 WHEN 'September' THEN 9
            WHEN 'October' THEN 10 WHEN 'November' THEN 11 WHEN 'December' THEN 12
        END DESC
    ''').fetchall()
    
    # Get employees list for filter
    employees = conn.execute('''
        SELECT DISTINCT e.emp_id, e.first_name, e.last_name
        FROM employees e
        JOIN work_progress wp ON e.emp_id = wp.emp_id
        ORDER BY e.first_name
    ''').fetchall()
    
    # Get projects list for filter
    projects = conn.execute('''
        SELECT DISTINCT p.proj_id, p.name
        FROM projects p
        JOIN work_progress wp ON p.proj_id = wp.proj_id
        ORDER BY p.name
    ''').fetchall()
    
    # Get clients list for filter
    clients = conn.execute('''
        SELECT DISTINCT p.client
        FROM projects p
        JOIN work_progress wp ON p.proj_id = wp.proj_id
        ORDER BY p.client
    ''').fetchall()
    
    # Calculate summary statistics
    total_entries = len(progress_data)
    if total_entries > 0:
        avg_percentage = sum(row['percentage'] for row in progress_data) / total_entries
        total_employees = len(set(row['emp_id'] for row in progress_data))
        total_projects = len(set(row['proj_id'] for row in progress_data))
    else:
        avg_percentage = 0
        total_employees = 0
        total_projects = 0
    
    conn.close()
    
    return render_template('admin_reports.html', 
                         reports=progress_data,
                         progress_data=progress_data,
                         months_years=months_years,
                         employees=employees,
                         projects=projects,
                         clients=clients,
                         total_entries=total_entries,
                         avg_percentage=avg_percentage,
                         total_employees=total_employees,
                         total_projects=total_projects)

# Report Generation Functions with Multiple Export Formats
@app.route('/admin/reports/export_csv')
def export_csv():
    if not session.get('is_admin'):
        return redirect(url_for('admin_login'))
    
    # Get filter parameters
    filter_month = request.args.get('month')
    filter_year = request.args.get('year')
    filter_employee = request.args.get('employee')
    filter_project = request.args.get('project')
    
    conn = get_db_connection()
    
    # Build query with filters
    query = '''
        SELECT wp.*, e.first_name, e.last_name, e.email, e.department, 
               p.name as project_name, p.client, p.description
        FROM work_progress wp
        JOIN employees e ON wp.emp_id = e.emp_id
        JOIN projects p ON wp.proj_id = p.proj_id
        WHERE wp.percentage > 0
    '''
    params = []
    
    if filter_month:
        query += ' AND wp.month = ?'
        params.append(filter_month)
    if filter_year:
        query += ' AND wp.year = ?'
        params.append(int(filter_year))
    if filter_employee:
        query += ' AND wp.emp_id = ?'
        params.append(filter_employee)
    if filter_project:
        query += ' AND wp.proj_id = ?'
        params.append(filter_project)
    
    query += ' ORDER BY wp.year DESC, wp.month, e.first_name, p.name'
    
    progress_data = conn.execute(query, params).fetchall()
    conn.close()
    
    # Create CSV
    output = StringIO()
    writer = csv.writer(output)
    
    # Write header
    writer.writerow([
        'Employee ID', 'Employee Name', 'Email', 'Department',
        'Project ID', 'Project Name', 'Client', 'Month', 'Year',
        'Percentage', 'Submission Date'
    ])
    
    # Write data
    for row in progress_data:
        writer.writerow([
            row['emp_id'],
            f"{row['first_name']} {row['last_name']}",
            row['email'],
            row['department'],
            row['proj_id'],
            row['project_name'],
            row['client'],
            row['month'],
            row['year'],
            row['percentage'],
            row['submission_date']
        ])
    
    output.seek(0)
    
    # Create response
    response = make_response(output.getvalue())
    response.headers['Content-Type'] = 'text/csv'
    
    # Generate filename
    filename = 'work_progress_report'
    if filter_month and filter_year:
        filename += f'_{filter_month}_{filter_year}'
    filename += f'_{datetime.now().strftime("%Y%m%d_%H%M%S")}.csv'
    
    response.headers['Content-Disposition'] = f'attachment; filename={filename}'
    
    flash('CSV Report exported successfully!', 'success')
    return response

@app.route('/admin/reports/export_excel')
def export_excel():
    if not session.get('is_admin'):
        return redirect(url_for('admin_login'))
    
    try:
        import pandas as pd
        from io import BytesIO
        
        # Get filter parameters
        filter_month = request.args.get('month')
        filter_year = request.args.get('year')
        filter_employee = request.args.get('employee')
        filter_project = request.args.get('project')
        
        conn = get_db_connection()
        
        # Build query with filters
        query = '''
            SELECT wp.emp_id as "Employee ID", 
                   e.first_name || ' ' || e.last_name as "Employee Name",
                   e.email as "Email", e.department as "Department",
                   wp.proj_id as "Project ID", p.name as "Project Name", 
                   p.client as "Client", wp.month as "Month", wp.year as "Year",
                   wp.percentage as "Percentage", wp.submission_date as "Submission Date"
            FROM work_progress wp
            JOIN employees e ON wp.emp_id = e.emp_id
            JOIN projects p ON wp.proj_id = p.proj_id
            WHERE wp.percentage > 0
        '''
        params = []
        
        if filter_month:
            query += ' AND wp.month = ?'
            params.append(filter_month)
        if filter_year:
            query += ' AND wp.year = ?'
            params.append(int(filter_year))
        if filter_employee:
            query += ' AND wp.emp_id = ?'
            params.append(filter_employee)
        if filter_project:
            query += ' AND wp.proj_id = ?'
            params.append(filter_project)
        
        query += ' ORDER BY wp.year DESC, wp.month, e.first_name, p.name'
        
        # Create DataFrame
        df = pd.read_sql_query(query, conn, params=params)
        conn.close()
        
        # Create Excel file
        output = BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='Work Progress', index=False)
            
            # Add summary sheet
            summary_data = {
                'Total Entries': [len(df)],
                'Average Percentage': [df['Percentage'].mean() if len(df) > 0 else 0],
                'Total Employees': [df['Employee ID'].nunique() if len(df) > 0 else 0],
                'Total Projects': [df['Project ID'].nunique() if len(df) > 0 else 0]
            }
            summary_df = pd.DataFrame(summary_data)
            summary_df.to_excel(writer, sheet_name='Summary', index=False)
        
        output.seek(0)
        
        # Create response
        response = make_response(output.getvalue())
        response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        
        # Generate filename
        filename = 'work_progress_report'
        if filter_month and filter_year:
            filename += f'_{filter_month}_{filter_year}'
        filename += f'_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx'
        
        response.headers['Content-Disposition'] = f'attachment; filename={filename}'
        
        flash('Excel Report exported successfully!', 'success')
        return response
        
    except ImportError:
        flash('Excel export requires pandas and openpyxl. Install with: pip install pandas openpyxl', 'error')
        return redirect(url_for('admin_reports'))
    except Exception as e:
        flash(f'Error creating Excel file: {str(e)}', 'error')
        return redirect(url_for('admin_reports'))

@app.route('/admin/reports/export_pdf')
def export_pdf():
    if not session.get('is_admin'):
        return redirect(url_for('admin_login'))
    
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from io import BytesIO
        
        # Get filter parameters
        filter_month = request.args.get('month')
        filter_year = request.args.get('year')
        filter_employee = request.args.get('employee')
        filter_project = request.args.get('project')
        
        conn = get_db_connection()
        
        # Build query with filters
        query = '''
            SELECT wp.*, e.first_name, e.last_name, e.email, e.department, 
                   p.name as project_name, p.client, p.description
            FROM work_progress wp
            JOIN employees e ON wp.emp_id = e.emp_id
            JOIN projects p ON wp.proj_id = p.proj_id
            WHERE wp.percentage > 0
        '''
        params = []
        
        if filter_month:
            query += ' AND wp.month = ?'
            params.append(filter_month)
        if filter_year:
            query += ' AND wp.year = ?'
            params.append(int(filter_year))
        if filter_employee:
            query += ' AND wp.emp_id = ?'
            params.append(filter_employee)
        if filter_project:
            query += ' AND wp.proj_id = ?'
            params.append(filter_project)
        
        query += ' ORDER BY wp.year DESC, wp.month, e.first_name, p.name'
        
        progress_data = conn.execute(query, params).fetchall()
        conn.close()
        
        # Create PDF
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        elements = []
        
        # Styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=30,
            textColor=colors.HexColor('#2c3e50')
        )
        
        # Title
        title = f"CAD Work Progress Report"
        if filter_month and filter_year:
            title += f" - {filter_month} {filter_year}"
        elements.append(Paragraph(title, title_style))
        elements.append(Spacer(1, 12))
        
        # Summary
        summary_text = f"""
        <b>Report Summary:</b><br/>
        Total Entries: {len(progress_data)}<br/>
        Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}<br/>
        """
        elements.append(Paragraph(summary_text, styles['Normal']))
        elements.append(Spacer(1, 20))
        
        # Data table
        if progress_data:
            data = [['Employee', 'Project', 'Client', 'Month/Year', 'Percentage']]
            for row in progress_data:
                data.append([
                    f"{row['first_name']} {row['last_name']}",
                    row['project_name'][:30] + '...' if len(row['project_name']) > 30 else row['project_name'],
                    row['client'],
                    f"{row['month']} {row['year']}",
                    f"{row['percentage']}%"
                ])
            
            table = Table(data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            elements.append(table)
        else:
            elements.append(Paragraph("No data available for the selected criteria.", styles['Normal']))
        
        doc.build(elements)
        buffer.seek(0)
        
        # Create response
        response = make_response(buffer.getvalue())
        response.headers['Content-Type'] = 'application/pdf'
        
        # Generate filename
        filename = 'work_progress_report'
        if filter_month and filter_year:
            filename += f'_{filter_month}_{filter_year}'
        filename += f'_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf'
        
        response.headers['Content-Disposition'] = f'attachment; filename={filename}'
        
        flash('PDF Report exported successfully!', 'success')
        return response
        
    except ImportError:
        flash('PDF export requires reportlab. Install with: pip install reportlab', 'error')
        return redirect(url_for('admin_reports'))
    except Exception as e:
        flash(f'Error creating PDF file: {str(e)}', 'error')
        return redirect(url_for('admin_reports'))

@app.route('/admin/reports/export_word')
def export_word():
    if not session.get('is_admin'):
        return redirect(url_for('admin_login'))
    
    try:
        from docx import Document
        from docx.shared import Inches
        from io import BytesIO
        
        # Get filter parameters
        filter_month = request.args.get('month')
        filter_year = request.args.get('year')
        filter_employee = request.args.get('employee')
        filter_project = request.args.get('project')
        
        conn = get_db_connection()
        
        # Build query with filters
        query = '''
            SELECT wp.*, e.first_name, e.last_name, e.email, e.department, 
                   p.name as project_name, p.client, p.description
            FROM work_progress wp
            JOIN employees e ON wp.emp_id = e.emp_id
            JOIN projects p ON wp.proj_id = p.proj_id
            WHERE wp.percentage > 0
        '''
        params = []
        
        if filter_month:
            query += ' AND wp.month = ?'
            params.append(filter_month)
        if filter_year:
            query += ' AND wp.year = ?'
            params.append(int(filter_year))
        if filter_employee:
            query += ' AND wp.emp_id = ?'
            params.append(filter_employee)
        if filter_project:
            query += ' AND wp.proj_id = ?'
            params.append(filter_project)
        
        query += ' ORDER BY wp.year DESC, wp.month, e.first_name, p.name'
        
        progress_data = conn.execute(query, params).fetchall()
        conn.close()
        
        # Create Word document
        doc = Document()
        
        # Title
        title = f"CAD Work Progress Report"
        if filter_month and filter_year:
            title += f" - {filter_month} {filter_year}"
        doc.add_heading(title, 0)
        
        # Summary
        summary = doc.add_paragraph()
        summary.add_run('Report Summary:\n').bold = True
        summary.add_run(f'Total Entries: {len(progress_data)}\n')
        summary.add_run(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}\n')
        
        # Data table
        if progress_data:
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            
            # Header
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Employee'
            hdr_cells[1].text = 'Project'
            hdr_cells[2].text = 'Client'
            hdr_cells[3].text = 'Month/Year'
            hdr_cells[4].text = 'Percentage'
            
            # Data rows
            for row in progress_data:
                row_cells = table.add_row().cells
                row_cells[0].text = f"{row['first_name']} {row['last_name']}"
                row_cells[1].text = row['project_name']
                row_cells[2].text = row['client']
                row_cells[3].text = f"{row['month']} {row['year']}"
                row_cells[4].text = f"{row['percentage']}%"
        else:
            doc.add_paragraph("No data available for the selected criteria.")
        
        # Save to BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Create response
        response = make_response(buffer.getvalue())
        response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        
        # Generate filename
        filename = 'work_progress_report'
        if filter_month and filter_year:
            filename += f'_{filter_month}_{filter_year}'
        filename += f'_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
        
        response.headers['Content-Disposition'] = f'attachment; filename={filename}'
        
        flash('Word Report exported successfully!', 'success')
        return response
        
    except ImportError:
        flash('Word export requires python-docx. Install with: pip install python-docx', 'error')
        return redirect(url_for('admin_reports'))
    except Exception as e:
        flash(f'Error creating Word file: {str(e)}', 'error')
        return redirect(url_for('admin_reports'))

@app.route('/admin/reports/summary')
def reports_summary():
    if not session.get('is_admin'):
        return redirect(url_for('admin_login'))
    
    conn = get_db_connection()
    
    # Get summary statistics
    summary_data = {
        'total_employees': conn.execute('SELECT COUNT(*) as count FROM employees WHERE is_admin = 0').fetchone()['count'],
        'total_projects': conn.execute('SELECT COUNT(*) as count FROM projects WHERE status = "Active"').fetchone()['count'],
        'total_submissions': conn.execute('SELECT COUNT(*) as count FROM work_progress').fetchone()['count'],
    }
    
    # Monthly completion rates
    monthly_stats = conn.execute('''
        SELECT month, year, 
               COUNT(DISTINCT emp_id) as employees_submitted,
               AVG(percentage) as avg_percentage,
               COUNT(*) as total_entries
        FROM work_progress 
        GROUP BY month, year 
        ORDER BY year DESC, 
        CASE month 
            WHEN 'January' THEN 1 WHEN 'February' THEN 2 WHEN 'March' THEN 3
            WHEN 'April' THEN 4 WHEN 'May' THEN 5 WHEN 'June' THEN 6
            WHEN 'July' THEN 7 WHEN 'August' THEN 8 WHEN 'September' THEN 9
            WHEN 'October' THEN 10 WHEN 'November' THEN 11 WHEN 'December' THEN 12
        END DESC
    ''').fetchall()
    
    # Employee performance summary
    employee_stats = conn.execute('''
        SELECT e.first_name, e.last_name, e.emp_id,
               COUNT(wp.id) as months_submitted,
               AVG(wp.percentage) as avg_percentage,
               MAX(wp.submission_date) as last_submission
        FROM employees e
        LEFT JOIN work_progress wp ON e.emp_id = wp.emp_id
        WHERE e.is_admin = 0
        GROUP BY e.emp_id, e.first_name, e.last_name
        ORDER BY months_submitted DESC, avg_percentage DESC
    ''').fetchall()
    
    # Project utilization
    project_stats = conn.execute('''
        SELECT p.name, p.proj_id, p.client,
               COUNT(wp.id) as times_worked,
               AVG(wp.percentage) as avg_allocation,
               COUNT(DISTINCT wp.emp_id) as unique_employees
        FROM projects p
        LEFT JOIN work_progress wp ON p.proj_id = wp.proj_id
        GROUP BY p.proj_id, p.name, p.client
        ORDER BY times_worked DESC
    ''').fetchall()
    
    conn.close()
    
    return jsonify({
        'summary': summary_data,
        'monthly_stats': [dict(row) for row in monthly_stats],
        'employee_stats': [dict(row) for row in employee_stats],
        'project_stats': [dict(row) for row in project_stats]
    })

# Database Management Functions
@app.route('/admin/database/backup')
def backup_database():
    if not session.get('is_admin'):
        return redirect(url_for('admin_login'))
    
    try:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        backup_filename = f'work_tracker_backup_{timestamp}.db'
        backup_path = os.path.join(os.getcwd(), 'backups', backup_filename)
        
        # Create backups directory if it doesn't exist
        os.makedirs(os.path.dirname(backup_path), exist_ok=True)
        
        # Copy database file
        shutil.copy2(DATABASE, backup_path)
        
        flash(f'Database backup created successfully: {backup_filename}', 'success')
    except Exception as e:
        flash(f'Error creating backup: {str(e)}', 'error')
    
    return redirect(url_for('admin_database'))

@app.route('/admin/database/clear_progress', methods=['POST'])
def clear_progress_data():
    if not session.get('is_admin'):
        return redirect(url_for('admin_login'))
    
    try:
        conn = get_db_connection()
        # Clear all work progress data
        deleted_count = conn.execute('DELETE FROM work_progress').rowcount
        conn.commit()
        conn.close()
        
        flash(f'Successfully cleared {deleted_count} progress entries', 'success')
    except Exception as e:
        flash(f'Error clearing progress data: {str(e)}', 'error')
    
    return redirect(url_for('admin_database'))

@app.route('/admin/database/wipe_all', methods=['POST'])
def wipe_database():
    if not session.get('is_admin'):
        return redirect(url_for('admin_login'))
    
    try:
        conn = get_db_connection()
        # Clear all tables except keep admin user
        conn.execute('DELETE FROM work_progress')
        conn.execute('DELETE FROM projects')
        conn.execute('DELETE FROM employees WHERE is_admin = 0')
        conn.commit()
        conn.close()
        
        flash('Database wiped successfully. Only admin user retained.', 'success')
    except Exception as e:
        flash(f'Error wiping database: {str(e)}', 'error')
    
    return redirect(url_for('admin_database'))

@app.route('/admin/database/restore/<filename>', methods=['POST'])
def restore_database(filename):
    if not session.get('is_admin'):
        return redirect(url_for('admin_login'))
    
    try:
        backup_path = os.path.join(os.getcwd(), 'backups', filename)
        
        if not os.path.exists(backup_path):
            flash('Backup file not found', 'error')
            return redirect(url_for('admin_database'))
        
        # Create backup of current database before restore
        current_backup = f'pre_restore_backup_{datetime.now().strftime("%Y%m%d_%H%M%S")}.db'
        current_backup_path = os.path.join(os.getcwd(), 'backups', current_backup)
        shutil.copy2(DATABASE, current_backup_path)
        
        # Restore from backup
        shutil.copy2(backup_path, DATABASE)
        
        flash(f'Database restored from {filename}. Current database backed up as {current_backup}', 'success')
    except Exception as e:
        flash(f'Error restoring database: {str(e)}', 'error')
    
    return redirect(url_for('admin_database'))

@app.route('/admin/database/delete_backup/<filename>', methods=['POST'])
def delete_backup(filename):
    if not session.get('is_admin'):
        return redirect(url_for('admin_login'))
    
    try:
        backup_path = os.path.join(os.getcwd(), 'backups', filename)
        
        if os.path.exists(backup_path):
            os.remove(backup_path)
            flash(f'Backup file {filename} deleted successfully', 'success')
        else:
            flash('Backup file not found', 'error')
    except Exception as e:
        flash(f'Error deleting backup: {str(e)}', 'error')
    
    return redirect(url_for('admin_database'))

@app.route('/admin/database/reset_sample_data', methods=['POST'])
def reset_sample_data():
    if not session.get('is_admin'):
        return redirect(url_for('admin_login'))
    
    try:
        # Create backup first
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        backup_filename = f'pre_reset_backup_{timestamp}.db'
        backup_path = os.path.join(os.getcwd(), 'backups', backup_filename)
        os.makedirs(os.path.dirname(backup_path), exist_ok=True)
        shutil.copy2(DATABASE, backup_path)
        
        # Remove current database
        if os.path.exists(DATABASE):
            os.remove(DATABASE)
        
        # Recreate with sample data
        init_database()
        
        flash(f'Sample data reset successfully. Previous data backed up as {backup_filename}', 'success')
    except Exception as e:
        flash(f'Error resetting sample data: {str(e)}', 'error')
    
    return redirect(url_for('admin_database'))

@app.route('/admin/database/integrity_check')
def integrity_check():
    if not session.get('is_admin'):
        return redirect(url_for('admin_login'))
    
    try:
        conn = get_db_connection()
        
        # Check for foreign key violations
        orphaned_progress = conn.execute('''
            SELECT wp.id, wp.emp_id, wp.proj_id, wp.month, wp.year
            FROM work_progress wp
            LEFT JOIN employees e ON wp.emp_id = e.emp_id
            LEFT JOIN projects p ON wp.proj_id = p.proj_id
            WHERE e.emp_id IS NULL OR p.proj_id IS NULL
        ''').fetchall()
        
        # Check for duplicate entries
        duplicates = conn.execute('''
            SELECT emp_id, proj_id, month, year, COUNT(*) as count
            FROM work_progress
            GROUP BY emp_id, proj_id, month, year
            HAVING COUNT(*) > 1
        ''').fetchall()
        
        # Check for invalid percentages
        invalid_percentages = conn.execute('''
            SELECT emp_id, month, year, SUM(percentage) as total
            FROM work_progress
            GROUP BY emp_id, month, year
            HAVING ABS(SUM(percentage) - 100.0) > 0.01
        ''').fetchall()
        
        conn.close()
        
        issues = []
        if orphaned_progress:
            issues.append(f'{len(orphaned_progress)} orphaned progress entries')
        if duplicates:
            issues.append(f'{len(duplicates)} duplicate entries')
        if invalid_percentages:
            issues.append(f'{len(invalid_percentages)} entries with invalid total percentages')
        
        if issues:
            flash(f'Database issues found: {", ".join(issues)}', 'warning')
        else:
            flash('Database integrity check passed. No issues found.', 'success')
            
    except Exception as e:
        flash(f'Error during integrity check: {str(e)}', 'error')
    
    return redirect(url_for('admin_database'))

@app.route('/admin/database/export_json')
def export_database_json():
    if not session.get('is_admin'):
        return redirect(url_for('admin_login'))
    
    try:
        conn = get_db_connection()
        
        # Export all data to JSON
        data = {
            'employees': [dict(row) for row in conn.execute('SELECT * FROM employees').fetchall()],
            'projects': [dict(row) for row in conn.execute('SELECT * FROM projects').fetchall()],
            'work_progress': [dict(row) for row in conn.execute('SELECT * FROM work_progress').fetchall()],
            'export_date': datetime.now().isoformat()
        }
        
        conn.close()
        
        # Create JSON response
        response = make_response(json.dumps(data, indent=2, default=str))
        response.headers['Content-Type'] = 'application/json'
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f'work_tracker_data_{timestamp}.json'
        response.headers['Content-Disposition'] = f'attachment; filename={filename}'
        
        flash('Database exported to JSON successfully!', 'success')
        return response
        
    except Exception as e:
        flash(f'Error exporting database: {str(e)}', 'error')
        return redirect(url_for('admin_database'))

@app.route('/admin/database/cleanup')
def cleanup_database():
    if not session.get('is_admin'):
        return redirect(url_for('admin_login'))
    
    try:
        conn = get_db_connection()
        
        # Remove duplicate entries (keep latest)
        conn.execute('''
            DELETE FROM work_progress 
            WHERE id NOT IN (
                SELECT MAX(id) FROM work_progress 
                GROUP BY emp_id, proj_id, month, year
            )
        ''')
        
        # Update foreign key constraints
        conn.execute('PRAGMA foreign_keys=OFF')
        conn.execute('VACUUM')
        conn.execute('PRAGMA foreign_keys=ON')
        
        conn.commit()
        conn.close()
        
        flash('Database cleanup completed successfully!', 'success')
    except Exception as e:
        flash(f'Error during cleanup: {str(e)}', 'error')
    
    return redirect(url_for('admin_database'))

@app.route('/admin/database/add_sample_progress', methods=['POST'])
def add_sample_progress():
    if not session.get('is_admin'):
        return redirect(url_for('admin_login'))
    
    try:
        conn = get_db_connection()
        
        # Add some sample progress data for July 2025
        sample_data = [
            ('EMP005', 'PROJ201', 'July', 2025, 25.0, '2025-08-01'),  # Subeesh
            ('EMP005', 'PROJ294', 'July', 2025, 30.0, '2025-08-01'),
            ('EMP005', 'PROJ286', 'July', 2025, 45.0, '2025-08-01'),
            ('EMP003', 'PROJ201', 'July', 2025, 20.0, '2025-08-02'),  # Ibrahim
            ('EMP003', 'PROJ208', 'July', 2025, 35.0, '2025-08-02'),
            ('EMP003', 'PROJ298', 'July', 2025, 45.0, '2025-08-02'),
            ('EMP001', 'PROJ297', 'July', 2025, 40.0, '2025-08-03'),  # Shaji
            ('EMP001', 'PROJ302', 'July', 2025, 30.0, '2025-08-03'),
            ('EMP001', 'PROJ303', 'July', 2025, 30.0, '2025-08-03'),
        ]
        
        for data in sample_data:
            conn.execute('''
                INSERT OR REPLACE INTO work_progress 
                (emp_id, proj_id, month, year, percentage, submission_date) 
                VALUES (?, ?, ?, ?, ?, ?)
            ''', data)
        
        conn.commit()
        conn.close()
        
        flash('Sample progress data added successfully!', 'success')
    except Exception as e:
        flash(f'Error adding sample data: {str(e)}', 'error')
    
    return redirect(url_for('admin_database'))

@app.route('/admin/database')
def admin_database():
    if not session.get('is_admin'):
        return redirect(url_for('admin_login'))
    
    conn = get_db_connection()
    
    # Get database statistics
    stats = {
        'total_employees': conn.execute('SELECT COUNT(*) as count FROM employees').fetchone()['count'],
        'total_projects': conn.execute('SELECT COUNT(*) as count FROM projects').fetchone()['count'],
        'total_progress_entries': conn.execute('SELECT COUNT(*) as count FROM work_progress').fetchone()['count'],
        'database_size': os.path.getsize(DATABASE) if os.path.exists(DATABASE) else 0
    }
    
    # Get recent activity
    recent_submissions = conn.execute('''
        SELECT wp.submission_date, e.first_name, e.last_name, wp.month, wp.year
        FROM work_progress wp
        JOIN employees e ON wp.emp_id = e.emp_id
        ORDER BY wp.submission_date DESC
        LIMIT 10
    ''').fetchall()
    
    # Check for data integrity issues
    orphaned_progress = conn.execute('''
        SELECT COUNT(*) as count FROM work_progress wp
        LEFT JOIN employees e ON wp.emp_id = e.emp_id
        LEFT JOIN projects p ON wp.proj_id = p.proj_id
        WHERE e.emp_id IS NULL OR p.proj_id IS NULL
    ''').fetchone()['count']
    
    # Get backup files if they exist
    backup_dir = os.path.join(os.getcwd(), 'backups')
    backup_files = []
    if os.path.exists(backup_dir):
        for f in os.listdir(backup_dir):
            if f.endswith('.db'):
                file_path = os.path.join(backup_dir, f)
                file_size = os.path.getsize(file_path)
                file_date = datetime.fromtimestamp(os.path.getmtime(file_path))
                backup_files.append({
                    'name': f,
                    'size': file_size,
                    'date': file_date
                })
        backup_files.sort(key=lambda x: x['date'], reverse=True)
    
    conn.close()
    
    return render_template('admin_database.html', 
                         stats=stats,
                         recent_submissions=recent_submissions,
                         orphaned_progress=orphaned_progress,
                         backup_files=backup_files)

# Enhanced Project Management Functions
@app.route('/admin/projects/bulk_update', methods=['POST'])
def bulk_update_projects():
    if not session.get('is_admin'):
        return redirect(url_for('admin_login'))
    
    try:
        action = request.form.get('action')
        selected_projects = request.form.getlist('selected_projects')
        
        if not selected_projects:
            flash('No projects selected', 'warning')
            return redirect(url_for('admin_projects'))
        
        conn = get_db_connection()
        
        if action == 'activate':
            for proj_id in selected_projects:
                conn.execute('UPDATE projects SET status = "Active" WHERE proj_id = ?', (proj_id,))
            flash(f'{len(selected_projects)} projects activated', 'success')
            
        elif action == 'deactivate':
            for proj_id in selected_projects:
                conn.execute('UPDATE projects SET status = "Inactive" WHERE proj_id = ?', (proj_id,))
            flash(f'{len(selected_projects)} projects deactivated', 'success')
            
        elif action == 'delete':
            deleted_count = 0
            for proj_id in selected_projects:
                # Check for progress data
                progress_count = conn.execute('SELECT COUNT(*) as count FROM work_progress WHERE proj_id = ?', (proj_id,)).fetchone()['count']
                if progress_count == 0:
                    conn.execute('DELETE FROM projects WHERE proj_id = ?', (proj_id,))
                    deleted_count += 1
                else:
                    conn.execute('UPDATE projects SET status = "Inactive" WHERE proj_id = ?', (proj_id,))
            
            if deleted_count > 0:
                flash(f'{deleted_count} projects deleted, others marked inactive (had progress data)', 'success')
            else:
                flash('Projects marked as inactive (all had progress data)', 'warning')
        
        conn.commit()
        conn.close()
        
    except Exception as e:
        flash(f'Error in bulk update: {str(e)}', 'error')
    
    return redirect(url_for('admin_projects'))

@app.route('/admin/projects/analytics')
def project_analytics():
    if not session.get('is_admin'):
        return redirect(url_for('admin_login'))
    
    conn = get_db_connection()
    
    # Project utilization analytics
    project_analytics = conn.execute('''
        SELECT p.*, 
               COUNT(wp.id) as total_allocations,
               AVG(wp.percentage) as avg_allocation,
               SUM(wp.percentage) as total_percentage,
               COUNT(DISTINCT wp.emp_id) as unique_employees,
               COUNT(DISTINCT wp.month || wp.year) as months_active
        FROM projects p
        LEFT JOIN work_progress wp ON p.proj_id = wp.proj_id
        GROUP BY p.proj_id
        ORDER BY total_percentage DESC
    ''').fetchall()
    
    # Monthly project trends
    monthly_trends = conn.execute('''
        SELECT wp.month, wp.year, p.name as project_name,
               SUM(wp.percentage) as total_allocation,
               COUNT(wp.emp_id) as employee_count
        FROM work_progress wp
        JOIN projects p ON wp.proj_id = p.proj_id
        GROUP BY wp.month, wp.year, p.proj_id, p.name
        ORDER BY wp.year DESC, 
        CASE wp.month 
            WHEN 'January' THEN 1 WHEN 'February' THEN 2 WHEN 'March' THEN 3
            WHEN 'April' THEN 4 WHEN 'May' THEN 5 WHEN 'June' THEN 6
            WHEN 'July' THEN 7 WHEN 'August' THEN 8 WHEN 'September' THEN 9
            WHEN 'October' THEN 10 WHEN 'November' THEN 11 WHEN 'December' THEN 12
        END DESC
    ''').fetchall()
    
    conn.close()
    
    return render_template('admin_project_analytics.html',
                         project_analytics=project_analytics,
                         monthly_trends=monthly_trends)

# Enhanced Employee Management Functions
@app.route('/admin/employees/bulk_update', methods=['POST'])
def bulk_update_employees():
    if not session.get('is_admin'):
        return redirect(url_for('admin_login'))
    
    try:
        action = request.form.get('action')
        selected_employees = request.form.getlist('selected_employees')
        
        if not selected_employees:
            flash('No employees selected', 'warning')
            return redirect(url_for('admin_employees'))
        
        conn = get_db_connection()
        
        if action == 'reset_passwords':
            for emp_id in selected_employees:
                employee = conn.execute('SELECT first_name FROM employees WHERE emp_id = ?', (emp_id,)).fetchone()
                if employee:
                    new_password = f"{employee['first_name'].lower()}@321"
                    hashed_password = generate_password_hash(new_password)
                    conn.execute('UPDATE employees SET password = ? WHERE emp_id = ?', (hashed_password, emp_id))
            flash(f'Passwords reset for {len(selected_employees)} employees', 'success')
            
        elif action == 'update_department':
            new_department = request.form.get('new_department')
            if new_department:
                for emp_id in selected_employees:
                    conn.execute('UPDATE employees SET department = ? WHERE emp_id = ?', (new_department, emp_id))
                flash(f'Department updated for {len(selected_employees)} employees', 'success')
            
        elif action == 'delete':
            deleted_count = 0
            for emp_id in selected_employees:
                # Check for progress data
                progress_count = conn.execute('SELECT COUNT(*) as count FROM work_progress WHERE emp_id = ?', (emp_id,)).fetchone()['count']
                if progress_count == 0:
                    conn.execute('DELETE FROM employees WHERE emp_id = ?', (emp_id,))
                    deleted_count += 1
            
            if deleted_count > 0:
                flash(f'{deleted_count} employees deleted, others not deleted (had progress data)', 'success')
            else:
                flash('No employees deleted (all had progress data)', 'warning')
        
        conn.commit()
        conn.close()
        
    except Exception as e:
        flash(f'Error in bulk update: {str(e)}', 'error')
    
    return redirect(url_for('admin_employees'))

@app.route('/admin/employees/analytics')
def employee_analytics():
    if not session.get('is_admin'):
        return redirect(url_for('admin_login'))
    
    conn = get_db_connection()
    
    # Employee performance analytics
    employee_analytics = conn.execute('''
        SELECT e.*, 
               COUNT(wp.id) as total_submissions,
               AVG(wp.percentage) as avg_workload,
               COUNT(DISTINCT wp.proj_id) as projects_worked,
               COUNT(DISTINCT wp.month || wp.year) as months_active,
               MAX(wp.submission_date) as last_submission,
               MIN(wp.submission_date) as first_submission
        FROM employees e
        LEFT JOIN work_progress wp ON e.emp_id = wp.emp_id
        WHERE e.is_admin = 0
        GROUP BY e.emp_id
        ORDER BY total_submissions DESC
    ''').fetchall()
    
    # Department summary
    department_summary = conn.execute('''
        SELECT e.department, 
               COUNT(e.emp_id) as employee_count,
               AVG(wp.percentage) as avg_workload,
               COUNT(wp.id) as total_submissions
        FROM employees e
        LEFT JOIN work_progress wp ON e.emp_id = wp.emp_id
        WHERE e.is_admin = 0
        GROUP BY e.department
        ORDER BY employee_count DESC
    ''').fetchall()
    
    conn.close()
    
    return render_template('admin_employee_analytics.html',
                         employee_analytics=employee_analytics,
                         department_summary=department_summary)

# Main execution block
if __name__ == '__main__':
    print("Initializing CAD Work Tracker...")
    init_database()
    print("Database initialized successfully!")
    print("Starting Flask application...")
    print("Access the application at:")
    print("  Local:   http://127.0.0.1:5000")
    print("  Network: http://0.0.0.0:5000")
    app.run(debug=True, host='0.0.0.0', port=5000)
